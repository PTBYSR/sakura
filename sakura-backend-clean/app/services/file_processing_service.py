"""
File processing service for extracting and chunking content from various file types.
Supports PDF, DOCX, TXT, CSV, XLSX files.
"""
import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from datetime import datetime
from pymongo.database import Database

from app.core.settings import Settings

logger = logging.getLogger(__name__)


class FileProcessingService:
    """Service for processing uploaded files and extracting chunks."""
    
    def __init__(self, settings: Settings, db: Optional[Database] = None):
        self.settings = settings
        self.db = db
        
        # Configuration
        self.chunk_size = int(os.getenv("FILE_CHUNK_SIZE", "200"))
        self.overlap = int(os.getenv("FILE_CHUNK_OVERLAP", "50"))
        
        # Initialize database indexes if database is available
        if self.db is not None:
            self._initialize_indexes()
    
    def _initialize_indexes(self):
        """Create indexes on MongoDB collections for better query performance."""
        try:
            # Indexes for files collection
            files_collection = self.db.files
            files_collection.create_index("file_id", unique=True)
            files_collection.create_index("dashboard_user_id")
            files_collection.create_index("status")
            files_collection.create_index([("dashboard_user_id", 1), ("status", 1)])
            logger.info("✅ Created indexes on files collection")
            
            # Indexes for file-chunks collection
            chunks_collection = self.db["file-chunks"]
            chunks_collection.create_index("file_id")
            chunks_collection.create_index("dashboard_user_id")
            chunks_collection.create_index([("dashboard_user_id", 1), ("file_id", 1)])
            chunks_collection.create_index([("file_id", 1), ("chunk_index", 1)])
            # Text index for full-text search on chunk text
            chunks_collection.create_index([("text", "text")])
            logger.info("✅ Created indexes on file-chunks collection")
        except Exception as e:
            logger.warning(f"⚠️  Could not create indexes: {e}")
    
    def chunk_text(self, text: str, chunk_size: Optional[int] = None, overlap: Optional[int] = None) -> List[Dict]:
        """Split text into overlapping word chunks."""
        chunk_size = chunk_size or self.chunk_size
        overlap = overlap or self.overlap
        
        words = re.findall(r"\S+", text)
        chunks = []
        i = 0
        while i < len(words):
            start = i
            end = min(i + chunk_size, len(words))
            chunk_words = words[start:end]
            chunk_text = " ".join(chunk_words)
            chunks.append({
                "id": f"chunk_{start}_{end}",
                "start_word": start,
                "end_word": end - 1,
                "text": chunk_text
            })
            i = i + chunk_size - overlap
        return chunks
    
    def extract_text_from_pdf(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from PDF file. Returns (text, page_count)."""
        try:
            import PyPDF2
            text_parts = []
            page_count = 0
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(text)
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
            
            full_text = "\n\n".join(text_parts)
            return full_text, page_count
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            raise Exception("PDF processing not available. Please install PyPDF2.")
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise
    
    def extract_text_from_docx(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from DOCX file. Returns (text, page_count)."""
        try:
            from docx import Document
            doc = Document(file_path)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = "\n\n".join(text_parts)
            # DOCX doesn't have page count, estimate based on content length
            estimated_pages = max(1, len(full_text) // 2000)  # Rough estimate: 2000 chars per page
            return full_text, estimated_pages
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise Exception("DOCX processing not available. Please install python-docx.")
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise
    
    def extract_text_from_txt(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from TXT file. Returns (text, page_count)."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            # Estimate pages (rough: 2000 chars per page)
            estimated_pages = max(1, len(text) // 2000)
            return text, estimated_pages
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            raise
    
    def extract_text_from_csv(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from CSV file. Returns (text, row_count)."""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to text representation
            text_parts = []
            # Add headers
            text_parts.append("Headers: " + ", ".join(df.columns.tolist()))
            # Add rows as text
            for idx, row in df.iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            row_count = len(df)
            return full_text, row_count
        except ImportError:
            logger.error("pandas not installed. Install with: pip install pandas")
            raise Exception("CSV processing not available. Please install pandas.")
        except Exception as e:
            logger.error(f"Error extracting text from CSV: {e}")
            raise
    
    def extract_text_from_xlsx(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from XLSX file. Returns (text, sheet_count)."""
        try:
            import pandas as pd
            excel_file = pd.ExcelFile(file_path)
            
            text_parts = []
            sheet_count = 0
            
            for sheet_name in excel_file.sheet_names:
                sheet_count += 1
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")
                text_parts.append("Headers: " + ", ".join(df.columns.tolist()))
                
                for idx, row in df.iterrows():
                    row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                    text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            return full_text, sheet_count
        except ImportError:
            logger.error("pandas and openpyxl not installed. Install with: pip install pandas openpyxl")
            raise Exception("XLSX processing not available. Please install pandas and openpyxl.")
        except Exception as e:
            logger.error(f"Error extracting text from XLSX: {e}")
            raise
    
    def extract_text_from_file(self, file_path: Path, file_type: str) -> Tuple[str, int]:
        """Extract text from file based on type. Returns (text, page_count)."""
        file_type_lower = file_type.lower()
        
        if file_type_lower == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_type_lower in ['docx', 'doc']:
            return self.extract_text_from_docx(file_path)
        elif file_type_lower == 'txt':
            return self.extract_text_from_txt(file_path)
        elif file_type_lower == 'csv':
            return self.extract_text_from_csv(file_path)
        elif file_type_lower in ['xlsx', 'xls']:
            return self.extract_text_from_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def process_file(
        self, 
        file_path: Path, 
        original_filename: str,
        file_type: str,
        file_id: str,
        dashboard_user_id: Optional[str] = None
    ) -> Dict:
        """Process a file: extract text, chunk it, and store in MongoDB."""
        now = datetime.now()
        
        # Get existing metadata from database or create new
        if self.db is not None:
            files_collection = self.db.files
            existing = files_collection.find_one({"file_id": file_id})
            
            if not existing:
                # Create new file entry in database
                file_doc = {
                    "file_id": file_id,
                    "name": original_filename,
                    "original_name": original_filename,
                    "type": file_type.lower(),
                    "status": "pending",
                    "pages_extracted": 0,
                    "total_chunks": 0,
                    "dashboard_user_id": dashboard_user_id,
                    "created_at": now,
                    "last_updated": now,
                    "error": None,
                    "config": {
                        "chunk_size": self.chunk_size,
                        "overlap": self.overlap
                    }
                }
                files_collection.insert_one(file_doc)
                logger.info(f"✅ Created new file entry in database: {file_id}")
        
        # Update status to processing
        if self.db is not None:
            files_collection = self.db.files
            files_collection.update_one(
                {"file_id": file_id},
                {
                    "$set": {
                        "status": "processing",
                        "last_updated": now
                    }
                }
            )
        
        try:
            # Extract text from file
            logger.info(f"Extracting text from {file_type} file: {original_filename}")
            text, page_count = self.extract_text_from_file(file_path, file_type)
            
            if not text or not text.strip():
                raise Exception("No text could be extracted from the file")
            
            # Chunk the text
            chunks = self.chunk_text(text, self.chunk_size, self.overlap)
            total_chunks = len(chunks)
            
            # Prepare chunks for database storage
            all_chunks = []
            for i, chunk in enumerate(chunks):
                chunk_doc = {
                    "file_id": file_id,
                    "dashboard_user_id": dashboard_user_id,
                    "chunk_index": i,
                    "chunk_id": chunk["id"],
                    "start_word": chunk["start_word"],
                    "end_word": chunk["end_word"],
                    "text": chunk["text"],
                    "source": original_filename,
                    "created_at": now
                }
                all_chunks.append(chunk_doc)
            
            # Store chunks in MongoDB (bulk insert for efficiency)
            if self.db is not None and all_chunks:
                chunks_collection = self.db["file-chunks"]
                # Delete existing chunks for this file first (in case of re-processing)
                chunks_collection.delete_many({"file_id": file_id})
                # Bulk insert all chunks
                chunks_collection.insert_many(all_chunks)
                logger.info(f"✅ Stored {len(all_chunks)} chunks in MongoDB for file {file_id}")
            
            # Update metadata with success
            if self.db is not None:
                files_collection.update_one(
                    {"file_id": file_id},
                    {
                        "$set": {
                            "status": "completed",
                            "pages_extracted": page_count,
                            "total_chunks": total_chunks,
                            "last_updated": datetime.now()
                        }
                    }
                )
                # Return updated file document
                updated = files_collection.find_one({"file_id": file_id})
                return {
                    "id": updated["file_id"],
                    "name": updated.get("name", original_filename),
                    "originalName": updated.get("original_name", original_filename),
                    "type": updated.get("type", file_type.lower()),
                    "size": updated.get("size", 0),
                    "status": updated["status"],
                    "pagesExtracted": updated.get("pages_extracted", 0),
                    "totalChunks": updated.get("total_chunks", 0),
                    "createdAt": updated["created_at"].isoformat() if isinstance(updated["created_at"], datetime) else updated["created_at"],
                    "lastUpdated": updated["last_updated"].isoformat() if isinstance(updated["last_updated"], datetime) else updated["last_updated"],
                    "error": updated.get("error"),
                    "dashboard_user_id": updated.get("dashboard_user_id")
                }
            else:
                # Fallback (no database)
                return {
                    "id": file_id,
                    "name": original_filename,
                    "originalName": original_filename,
                    "type": file_type.lower(),
                    "size": 0,
                    "status": "completed",
                    "pagesExtracted": page_count,
                    "totalChunks": total_chunks,
                    "createdAt": now.isoformat(),
                    "lastUpdated": now.isoformat(),
                    "error": None
                }
            
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Error processing file {original_filename}: {error_msg}")
            
            # Update metadata with error
            if self.db is not None:
                files_collection = self.db.files
                files_collection.update_one(
                    {"file_id": file_id},
                    {
                        "$set": {
                            "status": "error",
                            "error": error_msg,
                            "last_updated": datetime.now()
                        }
                    }
                )
            
            raise
    
    def get_file(self, file_id: str) -> Optional[Dict]:
        """Get file metadata by ID from database."""
        if self.db is not None:
            files_collection = self.db.files
            file = files_collection.find_one({"file_id": file_id})
            if file:
                return {
                    "id": file["file_id"],
                    "name": file.get("name", ""),
                    "originalName": file.get("original_name", ""),
                    "type": file.get("type", ""),
                    "size": file.get("size", 0),
                    "status": file["status"],
                    "pagesExtracted": file.get("pages_extracted", 0),
                    "totalChunks": file.get("total_chunks", 0),
                    "createdAt": file["created_at"].isoformat() if isinstance(file["created_at"], datetime) else file["created_at"],
                    "lastUpdated": file["last_updated"].isoformat() if isinstance(file["last_updated"], datetime) else file["last_updated"],
                    "error": file.get("error"),
                    "dashboard_user_id": file.get("dashboard_user_id")
                }
        return None
    
    def list_files(self, dashboard_user_id: Optional[str] = None) -> List[Dict]:
        """List all files, optionally filtered by dashboard_user_id."""
        if self.db is not None:
            files_collection = self.db.files
            query = {}
            if dashboard_user_id:
                query["dashboard_user_id"] = dashboard_user_id
            
            files = list(files_collection.find(query).sort("created_at", -1))
            # Convert MongoDB documents to dict format
            result = []
            for file in files:
                result.append({
                    "id": file["file_id"],
                    "name": file.get("name", ""),
                    "originalName": file.get("original_name", ""),
                    "type": file.get("type", ""),
                    "size": file.get("size", 0),
                    "status": file["status"],
                    "pagesExtracted": file.get("pages_extracted", 0),
                    "totalChunks": file.get("total_chunks", 0),
                    "createdAt": file["created_at"].isoformat() if isinstance(file["created_at"], datetime) else file["created_at"],
                    "lastUpdated": file["last_updated"].isoformat() if isinstance(file["last_updated"], datetime) else file["last_updated"],
                    "error": file.get("error"),
                    "dashboard_user_id": file.get("dashboard_user_id")
                })
            return result
        return []
    
    def delete_file(self, file_id: str) -> bool:
        """Delete a file and its chunks from database."""
        if self.db is not None:
            files_collection = self.db.files
            file = files_collection.find_one({"file_id": file_id})
            if not file:
                return False
            
            # Delete file metadata
            files_collection.delete_one({"file_id": file_id})
            
            # Delete all chunks for this file
            chunks_collection = self.db["file-chunks"]
            result = chunks_collection.delete_many({"file_id": file_id})
            logger.info(f"✅ Deleted {result.deleted_count} chunks from database for file {file_id}")
            
            return True
        return False
    
    def get_file_chunks(self, file_id: str) -> List[Dict]:
        """Get all chunks for a file from database."""
        if self.db is not None:
            chunks_collection = self.db["file-chunks"]
            chunks = list(chunks_collection.find(
                {"file_id": file_id}
            ).sort("chunk_index", 1))
            
            # Convert MongoDB documents to dict format (compatible with existing code)
            result = []
            for chunk in chunks:
                result.append({
                    "id": chunk.get("chunk_id", ""),
                    "start_word": chunk.get("start_word", 0),
                    "end_word": chunk.get("end_word", 0),
                    "text": chunk.get("text", ""),
                    "source": chunk.get("source", ""),
                    "chunk_index": chunk.get("chunk_index", 0),
                    "file_id": chunk.get("file_id", ""),
                })
            return result
        return []


# Global service instance
_file_processing_service: Optional[FileProcessingService] = None


def init_file_processing_service(settings: Settings, db: Optional[Database] = None) -> FileProcessingService:
    """Initialize file processing service with optional database."""
    global _file_processing_service
    _file_processing_service = FileProcessingService(settings, db=db)
    return _file_processing_service


def get_file_processing_service() -> FileProcessingService:
    """Get file processing service instance."""
    if not _file_processing_service:
        raise Exception("File processing service not initialized. Call init_file_processing_service() first.")
    return _file_processing_service

